#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 页面边距 ──
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)

def set_heading(para, text, size=16, bold=True, color=RGBColor(0x1F, 0x49, 0x7D)):
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_heading(p, text, size=18)
    elif level == 2:
        set_heading(p, text, size=14, color=RGBColor(0x2E, 0x74, 0xB5))
    else:
        set_heading(p, text, size=12, color=RGBColor(0x40, 0x40, 0x40))
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2E74B5')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'DEEAF1')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ══════════════════════════════════════════
# 封面标题
# ══════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_heading(title, "预测市场 GMGN — 市场战略文档", size=22, color=RGBColor(0x1F, 0x49, 0x7D))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("市场总监 OKR · GTM 策略 · 冷启动执行计划")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = date_p.add_run("2026年Q1")
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()

# ══════════════════════════════════════════
# 一、产品定位
# ══════════════════════════════════════════
add_heading(doc, "一、产品定位", 1)
p = doc.add_paragraph()
p.add_run("产品名称：").bold = True
p.add_run("预测市场 GMGN（聚合预测市场平台）")

p2 = doc.add_paragraph()
p2.add_run("核心功能：").bold = True
p2.add_run("聚合预测市场数据 · 聪明钱追踪 · 跟单功能")

p3 = doc.add_paragraph()
p3.add_run("核心定位：").bold = True
p3.add_run("第一个让用户能在预测市场跟单聪明钱的工具，降低预测市场信息不对称")

p4 = doc.add_paragraph()
p4.add_run("目标市场：").bold = True
p4.add_run("英文圈 — 现有预测市场用户（Polymarket/Kalshi）+ 加密交易用户")

doc.add_paragraph()

# ══════════════════════════════════════════
# 二、目标用户
# ══════════════════════════════════════════
add_heading(doc, "二、目标用户", 1)

add_heading(doc, "用户 A：预测市场原生用户（优先打）", 2)
doc.add_paragraph("• 平台：Polymarket、Kalshi、Manifold 用户")
doc.add_paragraph("• 痛点：不知道聪明钱在押什么，自己判断费力")
doc.add_paragraph("• 优势：已懂产品逻辑，转化快")

add_heading(doc, "用户 B：加密/链上交易用户（放大用）", 2)
doc.add_paragraph("• 平台：GMGN、DEX Screener、链上钱包追踪用户")
doc.add_paragraph("• 痛点：习惯跟单但缺少预测市场入口")
doc.add_paragraph("• 优势：圈子大，跟单心智成熟")

doc.add_paragraph()

# ══════════════════════════════════════════
# 三、GTM 市场策略（三阶段）
# ══════════════════════════════════════════
add_heading(doc, "三、GTM 市场策略（三阶段）", 1)

headers = ["阶段", "时间", "核心目标", "主要动作", "成功标志"]
rows = [
    ["阶段一\n冷启动", "第1-2个月", "建立品牌认知\n积累种子用户",
     "• 建Twitter内容账号\n• 进入Polymarket Discord/Reddit\n• 每日发Smart Money Report",
     "Twitter 500+粉\nWaitlist 200人"],
    ["阶段二\n增长", "第3个月", "KOL背书\n产品上线",
     "• 签约3-5个中腰部KOL\n• 开放内测\n• 用户跟单战绩卡传播",
     "注册用户1000人\n活跃用户200人"],
    ["阶段三\n规模化", "Q2起", "品牌升级\n数据基础设施",
     "• 每周聪明钱报告\n• 加密媒体合作\n• 数据API对外开放",
     "媒体引用2+次\n自然增长启动"],
]
add_table(doc, headers, rows, col_widths=[0.8, 0.8, 1.2, 2.5, 1.3])

doc.add_paragraph()

# ══════════════════════════════════════════
# 四、冷启动渠道执行计划
# ══════════════════════════════════════════
add_heading(doc, "四、冷启动渠道执行计划", 1)

headers2 = ["渠道", "平台", "具体动作", "频次", "目标"]
rows2 = [
    ["Twitter/X\n主阵地", "Twitter/X",
     "每日发 Smart Money Report\n主动回复Polymarket官推和KOL\n格式统一，数据可视化",
     "每日1条", "5000粉\n互动率>3%"],
    ["社区渗透", "Polymarket Discord\nReddit r/predictionmarkets",
     "发聪明钱数据洞察\n回答用户问题\n不直接推产品，先建信任",
     "每周3条", "500+ upvotes\n被圈子认可"],
    ["KOL合作", "Twitter/X",
     "找1-10万粉中腰部KOL\n给独家早期访问权限\n提供专属数据dashboard",
     "签约3-5个", "首批真实用户\n战绩内容传播"],
    ["私域社群", "Telegram/Discord",
     "建立核心用户群\n每日推送聪明钱数据\n内测资格优先给群内用户",
     "每日维护", "200核心用户"],
]
add_table(doc, headers2, rows2, col_widths=[1.0, 1.2, 2.3, 0.8, 1.3])

doc.add_paragraph()

# ══════════════════════════════════════════
# 五、OKR（Q1）— 优化版（结合获客策略图）
# ══════════════════════════════════════════
add_heading(doc, "五、市场总监 OKR（Q1）— 优化版", 1)

# 策略背景说明
note = doc.add_paragraph()
note.add_run("策略背景：").bold = True
note.add_run("获客分两路——代理渠道（30%）+ 直客渠道（70%）。前期无手续费，BD代理难吸引，代理渠道以KOL激励替代。目标用户细分为加密货币用户和美股用户，地区聚焦欧美+东南亚。")
note.runs[1].font.color.rgb = RGBColor(0x50, 0x50, 0x50)
doc.add_paragraph()

okr_data = [
    ("O1", "建立品牌认知，覆盖欧美+东南亚核心用户圈"),
    ("O2", "代理渠道（30%）— 用KOL激励替代BD冷启动"),
    ("O3", "直客渠道（70%）— 多路并发获取真实用户"),
    ("O4", "产品上线首月，跑通两类用户转化路径"),
]

kr_data = {
    "O1": [
        ("KR1", "加密货币渠道：在欧美+东南亚加密媒体发布 8 篇品牌内容，累计曝光 >50万", "第1季度末"),
        ("KR2", "美股渠道：在财经类媒体/账号发布 4 篇股票预测内容，覆盖欧美受众", "第1季度末"),
        ("KR3", "Twitter 账号粉丝达到 5,000，每条内容互动率 >3%", "第1季度末"),
    ],
    "O2": [
        ("KR1", "签约 5-8 个欧美+东南亚加密/美股中腰部KOL（1-10万粉），提供激励+定制化活动支持", "第2个月末"),
        ("KR2", "KOL合作产出内容 >20 条，带来 Waitlist 注册用户 >300", "第2个月末"),
        ("KR3", "举办至少 2 场预测体验定制化活动，每场参与人数 >200", "第1季度末"),
    ],
    "O3": [
        ("KR1（积分奖励）", "上线交易/邀请/持仓积分体系，首月参与用户 >500，邀请裂变带来新用户 >200", "上线后第1个月"),
        ("KR2（全球大使）", "招募全球大使 20 人（欧美+东南亚），瓜分奖池激励，大使带来注册用户 >300", "第1季度末"),
        ("KR3（SEO）", "完成 20 个核心关键词布局（prediction market、Polymarket alternative等），自然搜索流量 >5,000/月", "第1季度末"),
        ("KR4（邮件营销）", "建立流失用户召回机制，首月邮件打开率 >25%，召回率 >10%", "上线后第1个月"),
    ],
    "O4": [
        ("KR1", "加密货币用户首月注册 700，美股用户注册 300，合计 1,000", "上线后第1个月"),
        ("KR2", "跟单功能首月活跃用户 200，人均跟单次数 >3次", "上线后第1个月"),
        ("KR3", "用户自发分享跟单战绩内容 >50 条（战绩卡设计驱动）", "上线后第1个月"),
    ],
}

for obj_id, obj_text in okr_data:
    add_heading(doc, f"{obj_id}：{obj_text}", 2)
    headers3 = ["KR", "关键结果", "截止时间", "状态"]
    rows3 = [(kr, desc, deadline, "未开始") for kr, desc, deadline in kr_data[obj_id]]
    add_table(doc, headers3, rows3, col_widths=[1.2, 3.3, 1.3, 0.8])
    doc.add_paragraph()

# OKR 对比说明
add_heading(doc, "OKR 优化前后对比", 2)
headers_cmp = ["维度", "优化前", "优化后"]
rows_cmp = [
    ["用户分类", "未细分", "加密用户 + 美股用户分开打"],
    ["地区策略", "无", "聚焦欧美 + 东南亚"],
    ["代理渠道", "仅提KOL合作", "明确KOL激励替代BD，前期无手续费针对性方案"],
    ["直客增长机制", "无", "积分体系 + 全球大使 + SEO + 邮件营销四路并发"],
    ["用户召回", "无", "邮件营销覆盖流失用户名单"],
]
add_table(doc, headers_cmp, rows_cmp, col_widths=[1.2, 2.0, 3.4])
doc.add_paragraph()

# ══════════════════════════════════════════
# 六、所需资源
# ══════════════════════════════════════════
add_heading(doc, "六、启动所需核心资源", 1)

headers4 = ["资源", "用途", "优先级"]
rows4 = [
    ["每日聪明钱数据", "支撑 Twitter 日更内容 + 深度报告", "★★★ 最高"],
    ["英文内容运营", "执行日更 SOP，维护社区", "★★★ 最高"],
    ["KOL 合作预算", "早期背书，首批用户获取", "★★ 重要"],
    ["Waitlist 落地页", "蓄水种子用户", "★★ 重要"],
    ["私域社群（Telegram）", "核心用户运营 + 内测资格管理", "★ 可后续补充"],
]
add_table(doc, headers4, rows4, col_widths=[1.8, 3.5, 1.3])

doc.add_paragraph()
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_end = p_end.add_run("— END —")
run_end.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run_end.font.size = Pt(10)

doc.save("/Users/coco/agent-twitter/GMGN_市场战略文档_Q1.docx")
print("文档已生成：GMGN_市场战略文档_Q1.docx")
